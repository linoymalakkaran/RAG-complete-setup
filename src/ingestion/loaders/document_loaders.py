"""
Document loaders for various file formats.
Supports PDF, Word, text, images (OCR), and video transcripts.
"""

import io
from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
import pypdf
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import cv2
import numpy as np

from src.utils.logging_config import setup_logging

logger = setup_logging("rag.loaders")


class DocumentLoader(ABC):
    """
    Abstract base class for document loaders.
    All loaders should inherit from this and implement the load method.
    """
    
    @abstractmethod
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load document and extract content with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing:
                - content: str - Extracted text content
                - metadata: dict - Document metadata (author, date, etc.)
                - images: list - Extracted images (if any)
        """
        pass
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic metadata from file"""
        path = Path(file_path)
        return {
            'filename': path.name,
            'file_type': path.suffix,
            'file_size': path.stat().st_size,
            'modified_date': path.stat().st_mtime
        }


class PDFLoader(DocumentLoader):
    """
    PDF document loader with OCR fallback.
    
    This loader first tries to extract text directly from PDF.
    If no text is found (scanned PDF), it falls back to OCR.
    
    Features:
        - Direct text extraction
        - OCR fallback for scanned documents
        - Image extraction
        - Metadata extraction (author, title, etc.)
    """
    
    def __init__(self, use_ocr: bool = True, ocr_language: str = 'eng'):
        """
        Initialize PDF loader.
        
        Args:
            use_ocr: Whether to use OCR for scanned PDFs
            ocr_language: Language for OCR (default: English)
        """
        self.use_ocr = use_ocr
        self.ocr_language = ocr_language
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load PDF document"""
        logger.info(f"Loading PDF: {file_path}")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                metadata = self.extract_metadata(file_path)
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'num_pages': len(pdf_reader.pages)
                    })
                
                # Extract text from all pages
                content_parts = []
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(f"[Page {page_num}]\n{text}")
                
                content = "\n\n".join(content_parts)
                
                # If no text extracted and OCR is enabled, use OCR
                if not content.strip() and self.use_ocr:
                    logger.info(f"No text found in PDF, using OCR: {file_path}")
                    content = self._ocr_pdf(file_path)
                
                return {
                    'content': content,
                    'metadata': metadata,
                    'images': []  # TODO: Extract images if needed
                }
                
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    def _ocr_pdf(self, file_path: str) -> str:
        """
        Perform OCR on PDF pages.
        Converts PDF to images and runs OCR on each page.
        """
        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path)
            
            content_parts = []
            for page_num, image in enumerate(images, 1):
                # Run OCR on each page
                text = pytesseract.image_to_string(image, lang=self.ocr_language)
                if text.strip():
                    content_parts.append(f"[Page {page_num} - OCR]\n{text}")
            
            return "\n\n".join(content_parts)
            
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {str(e)}")
            return ""


class WordLoader(DocumentLoader):
    """
    Microsoft Word document loader (.docx).
    
    Extracts text, tables, and metadata from Word documents.
    """
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load Word document"""
        logger.info(f"Loading Word document: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = self.extract_metadata(file_path)
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            })
            
            # Extract text from paragraphs
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    content_parts.append(table_text)
            
            content = "\n\n".join(content_parts)
            
            return {
                'content': content,
                'metadata': metadata,
                'images': []
            }
            
        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {str(e)}")
            raise
    
    def _extract_table(self, table) -> str:
        """Extract text from Word table"""
        rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                rows.append(row_text)
        return "\n".join(rows)


class TextLoader(DocumentLoader):
    """
    Plain text and markdown file loader.
    Supports .txt, .md, .csv files.
    """
    
    def __init__(self, encoding: str = 'utf-8'):
        self.encoding = encoding
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load text file"""
        logger.info(f"Loading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding=self.encoding) as f:
                content = f.read()
            
            metadata = self.extract_metadata(file_path)
            
            return {
                'content': content,
                'metadata': metadata,
                'images': []
            }
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            raise


class ImageLoader(DocumentLoader):
    """
    Image loader with OCR support.
    Supports .png, .jpg, .jpeg files.
    
    Use cases:
        - Scanned documents
        - Screenshots
        - Diagrams with text
        - Product photos with labels
    """
    
    def __init__(self, ocr_language: str = 'eng', preprocess: bool = True):
        """
        Initialize image loader.
        
        Args:
            ocr_language: Language for OCR
            preprocess: Whether to preprocess image for better OCR
        """
        self.ocr_language = ocr_language
        self.preprocess = preprocess
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load and OCR image"""
        logger.info(f"Loading image: {file_path}")
        
        try:
            # Load image
            image = Image.open(file_path)
            
            # Preprocess if enabled
            if self.preprocess:
                image = self._preprocess_image(image)
            
            # Run OCR
            content = pytesseract.image_to_string(image, lang=self.ocr_language)
            
            metadata = self.extract_metadata(file_path)
            metadata.update({
                'width': image.width,
                'height': image.height,
                'mode': image.mode
            })
            
            return {
                'content': content,
                'metadata': metadata,
                'images': [file_path]  # Store reference to original image
            }
            
        except Exception as e:
            logger.error(f"Error loading image {file_path}: {str(e)}")
            raise
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image for better OCR results.
        
        Techniques:
            - Convert to grayscale
            - Increase contrast
            - Denoise
            - Binarization
        """
        # Convert PIL to OpenCV format
        img_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        
        # Convert to grayscale
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        
        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Adaptive thresholding for binarization
        binary = cv2.adaptiveThreshold(
            denoised, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        
        # Convert back to PIL
        return Image.fromarray(binary)


class VideoTranscriptLoader(DocumentLoader):
    """
    Video transcript loader.
    
    Can extract transcripts from:
        - Local video files (using Whisper)
        - YouTube URLs (using youtube-transcript-api)
    """
    
    def __init__(self, use_whisper: bool = True):
        self.use_whisper = use_whisper
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load video and extract transcript.
        
        Note: This is a simplified version. Full implementation would use
        libraries like whisper or youtube-transcript-api.
        """
        logger.info(f"Loading video transcript: {file_path}")
        
        # Placeholder implementation
        # In production, you would use:
        # - whisper.load_model() and whisper.transcribe() for local videos
        # - YouTubeTranscriptApi for YouTube videos
        
        metadata = self.extract_metadata(file_path)
        
        return {
            'content': "Video transcript extraction requires additional setup.",
            'metadata': metadata,
            'images': []
        }


class LoaderFactory:
    """
    Factory class to create appropriate loader based on file type.
    
    Usage:
        >>> loader = LoaderFactory.get_loader("document.pdf")
        >>> result = loader.load("document.pdf")
    """
    
    _loaders = {
        '.pdf': PDFLoader,
        '.docx': WordLoader,
        '.txt': TextLoader,
        '.md': TextLoader,
        '.csv': TextLoader,
        '.png': ImageLoader,
        '.jpg': ImageLoader,
        '.jpeg': ImageLoader,
        '.mp4': VideoTranscriptLoader,
        '.avi': VideoTranscriptLoader
    }
    
    @classmethod
    def get_loader(cls, file_path: str, **kwargs) -> DocumentLoader:
        """
        Get appropriate loader for file type.
        
        Args:
            file_path: Path to file
            **kwargs: Additional arguments for loader
            
        Returns:
            DocumentLoader instance
            
        Raises:
            ValueError: If file type is not supported
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in cls._loaders:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        loader_class = cls._loaders[file_ext]
        return loader_class(**kwargs)
    
    @classmethod
    def supported_formats(cls) -> List[str]:
        """Get list of supported file formats"""
        return list(cls._loaders.keys())


def load_document(file_path: str, **kwargs) -> Dict[str, Any]:
    """
    Convenience function to load any supported document type.
    
    Args:
        file_path: Path to document
        **kwargs: Additional arguments for specific loader
        
    Returns:
        Document content and metadata
        
    Example:
        >>> doc = load_document("policies/leave_policy.pdf")
        >>> print(doc['content'][:100])
        >>> print(doc['metadata']['num_pages'])
    """
    loader = LoaderFactory.get_loader(file_path, **kwargs)
    return loader.load(file_path)
